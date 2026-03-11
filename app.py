import streamlit as st
import pandas as pd
import json
import uuid
import os
import plotly.express as px
import plotly.graph_objects as go
import docx
from fpdf import FPDF
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
import sqlite3
from data_generator import generate_resumes

# Load environment variables (Requirement 7)
load_dotenv()
import matplotlib.pyplot as plt
import numpy as np
from agent import ScreeningAgent
from observability import JobCriterion, State
from tools import ResumeParserTool

# Page Config
st.set_page_config(page_title="Resume Screening Assistant", layout="wide")

# Initialize Session State
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "username" not in st.session_state:
    st.session_state.username = None
if "auth_mode" not in st.session_state:
    st.session_state.auth_mode = "signup" # Default to Sign Up as requested

# Database Initialization (Requirement: Professional Tech Stack)
def init_db():
    conn = sqlite3.connect("users.db")
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users 
                 (username TEXT PRIMARY KEY, email TEXT UNIQUE, password TEXT)''')
    
    # Migrate from JSON if it exists and DB is empty
    c.execute("SELECT COUNT(*) FROM users")
    if c.fetchone()[0] == 0 and os.path.exists("users.json"):
        try:
            with open("users.json", "r") as f:
                old_users = json.load(f)
                for u in old_users:
                    c.execute("INSERT OR IGNORE INTO users (username, email, password) VALUES (?, ?, ?)",
                              (u["username"], u["email"], u["password"]))
        except: pass
    
    conn.commit()
    conn.close()

init_db()

def login_user(u_or_e, p):
    conn = sqlite3.connect("users.db")
    c = conn.cursor()
    c.execute("SELECT username FROM users WHERE (username=? OR email=?) AND password=?", (u_or_e, u_or_e, p))
    user = c.fetchone()
    conn.close()
    return user[0] if user else None

def register_user(u, e, p):
    try:
        conn = sqlite3.connect("users.db")
        c = conn.cursor()
        c.execute("INSERT INTO users (username, email, password) VALUES (?, ?, ?)", (u, e, p))
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        return False

def login_page():
    st.title("🔐 Resume Screening Assistant")
    
    if st.session_state.auth_mode == "login":
        st.subheader("Login to your account")
        u_or_e = st.text_input("Username or Email", key="login_u")
        p = st.text_input("Password", type="password", key="login_p")
        
        col1, col2 = st.columns([1, 2])
        if col1.button("Login", use_container_width=True):
            username = login_user(u_or_e, p)
            if username:
                st.session_state.authenticated = True
                st.session_state.username = username
                st.success("Logged in successfully!")
                st.rerun()
            else:
                st.error("Invalid credentials")
        
        if st.button("New user? Create an account (Sign Up)"):
            st.session_state.auth_mode = "signup"
            st.rerun()
            
    else:
        st.subheader("Create a new account")
        new_u = st.text_input("Choose Username", key="signup_u")
        new_e = st.text_input("Enter Email", key="signup_e")
        new_p = st.text_input("Choose Password", type="password", key="signup_p")
        confirm_p = st.text_input("Confirm Password", type="password", key="signup_p_confirm")
        
        col1, col2 = st.columns([1, 2])
        if col1.button("Register & Sign Up", use_container_width=True):
            if not new_u or not new_p or not new_e:
                st.error("Please fill all fields")
            elif new_p != confirm_p:
                st.error("Passwords don't match")
            elif "@" not in new_e or "." not in new_e:
                st.error("Please enter a valid email address")
            else:
                if register_user(new_u, new_e, new_p):
                    st.success("Account created successfully! Now you can Login.")
                    st.session_state.auth_mode = "login"
                    st.rerun()
                else:
                    st.error("Username or Email already exists")
        
        if st.button("Already have an account? Go to Login"):
            st.session_state.auth_mode = "login"
            st.rerun()

# Main App Guard
if not st.session_state.authenticated:
    login_page()
    st.stop()

# Sidebar - User Info & Logout
with st.sidebar:
    st.markdown(f"### 👤 Welcome, {st.session_state.username}")
    if st.button("🚪 Logout", use_container_width=True):
        st.session_state.authenticated = False
        st.session_state.username = None
        st.rerun()
    st.markdown("---")

if "run_id" not in st.session_state:
    st.session_state.run_id = None
if "results" not in st.session_state:
    st.session_state.results = []
if "metrics" not in st.session_state:
    st.session_state.metrics = {}
if "feedback" not in st.session_state:
    st.session_state.feedback = {}

# Sidebar - Configuration
st.sidebar.header("Run Configuration")
source_mode = st.sidebar.radio("Data Source", ["Synthetic Resumes", "Upload PDFs"])
scenarios = [
    "Custom", "Data Scientist", "Frontend Developer", "DevOps Engineer", 
    "Backend Developer", "ML Engineer", "Full Stack Developer", 
    "Data Engineer", "QA Automation", "Cybersecurity Analyst", "Cloud Architect"
]
scenario_type = st.sidebar.selectbox("Select Scenario (10+ Presets)", scenarios)

# Auto-clear results if scenario or source changes
if "prev_scenario" not in st.session_state:
    st.session_state.prev_scenario = scenario_type
if "prev_source" not in st.session_state:
    st.session_state.prev_source = source_mode

if st.session_state.prev_scenario != scenario_type or st.session_state.prev_source != source_mode:
    st.session_state.results = []
    st.session_state.metrics = {}
    st.session_state.prev_scenario = scenario_type
    st.session_state.prev_source = source_mode

seed = st.sidebar.number_input("Seed (for reproducibility)", value=42, step=1)
num_resumes = st.sidebar.slider("Number of Resumes", 5, 50, 20)
simulate_failure = st.sidebar.checkbox("Simulate Failure Scenario")

st.sidebar.subheader("Job Criteria & Weights")

# Preset Scenarios Logic
if scenario_type == "Data Scientist":
    skills_weight = 0.6; exp_weight = 0.2; edu_weight = 0.2
    req_skills = ["Python", "Machine Learning", "SQL", "Pandas"]
    min_exp = 3; edu_keywords = ["Data Science", "Mathematics", "Statistics"]
elif scenario_type == "Frontend Developer":
    skills_weight = 0.7; exp_weight = 0.2; edu_weight = 0.1
    req_skills = ["React", "CSS", "JavaScript", "HTML"]
    min_exp = 2; edu_keywords = ["Computer Science", "Information Technology"]
elif scenario_type == "DevOps Engineer":
    skills_weight = 0.5; exp_weight = 0.4; edu_weight = 0.1
    req_skills = ["Docker", "Kubernetes", "AWS", "Git"]
    min_exp = 5; edu_keywords = ["Engineering", "Computer Science"]
elif scenario_type == "Backend Developer":
    skills_weight = 0.6; exp_weight = 0.3; edu_weight = 0.1
    req_skills = ["Python", "Django", "SQL", "Redis"]
    min_exp = 4; edu_keywords = ["Computer Science", "Engineering"]
elif scenario_type == "ML Engineer":
    skills_weight = 0.7; exp_weight = 0.1; edu_weight = 0.2
    req_skills = ["PyTorch", "TensorFlow", "Scikit-Learn", "Python"]
    min_exp = 2; edu_keywords = ["Mathematics", "Statistics", "Data Science"]
elif scenario_type == "Full Stack Developer":
    skills_weight = 0.5; exp_weight = 0.4; edu_weight = 0.1
    req_skills = ["React", "Node.js", "MongoDB", "JavaScript"]
    min_exp = 3; edu_keywords = ["Computer Science", "IT"]
elif scenario_type == "Data Engineer":
    skills_weight = 0.6; exp_weight = 0.3; edu_weight = 0.1
    req_skills = ["Spark", "Hadoop", "SQL", "Airflow"]
    min_exp = 4; edu_keywords = ["Computer Science", "Data Science"]
elif scenario_type == "QA Automation":
    skills_weight = 0.5; exp_weight = 0.4; edu_weight = 0.1
    req_skills = ["Selenium", "Python", "JUnit", "TestNG"]
    min_exp = 2; edu_keywords = ["Engineering", "Computer Science"]
elif scenario_type == "Cybersecurity Analyst":
    skills_weight = 0.7; exp_weight = 0.2; edu_weight = 0.1
    req_skills = ["Nmap", "Metasploit", "Wireshark", "Python"]
    min_exp = 5; edu_keywords = ["Cybersecurity", "IT"]
elif scenario_type == "Cloud Architect":
    skills_weight = 0.4; exp_weight = 0.5; edu_weight = 0.1
    req_skills = ["AWS", "Azure", "Terraform", "CloudFormation"]
    min_exp = 8; edu_keywords = ["Engineering", "Computer Science"]
else:
    skills_weight = st.sidebar.slider("Skills Weight", 0.0, 1.0, 0.5)
    exp_weight = st.sidebar.slider("Experience Weight", 0.0, 1.0, 0.3)
    edu_weight = st.sidebar.slider("Education Weight", 0.0, 1.0, 0.2)
    req_skills = ["Python", "SQL", "FastAPI"]
    min_exp = 5
    edu_keywords = ["Computer Science", "Engineering"]

# Ensure weights sum to 1
total_w = skills_weight + exp_weight + edu_weight
if total_w > 0:
    skills_weight /= total_w; exp_weight /= total_w; edu_weight /= total_w

criteria = [
    JobCriterion(name="Skills Match", weight=skills_weight, required_skills=req_skills),
    JobCriterion(name="Experience Match", weight=exp_weight, min_experience=min_exp),
    JobCriterion(name="Education Match", weight=edu_weight, education_keywords=edu_keywords)
]

# Ensure upload directory exists on server
if not os.path.exists("uploaded_resumes"):
    os.makedirs("uploaded_resumes")

# Main UI
st.title("🚀 AI Resume Screening Assistant")
st.markdown("---")

uploaded_files = []
if source_mode == "Upload PDFs":
    uploaded_files = st.file_uploader("Upload Resumes (PDF)", type=["pdf"], accept_multiple_files=True)
    if uploaded_files:
        st.success(f"{len(uploaded_files)} PDF(s) uploaded successfully.")

col1, col2 = st.columns([1, 2])

with col1:
    st.subheader("Controls")
    if st.button("🚀 Start Screening Run", use_container_width=True):
        # Reset previous run logs
        if os.path.exists("runs.jsonl"):
            os.remove("runs.jsonl")
            
        st.session_state.run_id = str(uuid.uuid4())
        
        with st.spinner("Agent working..."):
            # 1. Data Source Handling
            resumes = []
            if source_mode == "Synthetic Resumes":
                resumes = generate_resumes(num_resumes=num_resumes, seed=seed)
            else:
                if not uploaded_files:
                    st.error("Please upload some PDFs first!")
                    st.stop()
                
                parser_tool = ResumeParserTool()
                for i, pdf in enumerate(uploaded_files):
                    try:
                        # Save file to disk for persistence
                        save_path = os.path.join("uploaded_resumes", pdf.name)
                        with open(save_path, "wb") as f:
                            f.write(pdf.getbuffer())
                            
                        parsed_resume = parser_tool.parse_pdf(pdf)
                        parsed_resume.id = i + 1 # Assign sequential IDs
                        resumes.append(parsed_resume)
                        st.sidebar.caption(f"✅ Saved: {pdf.name}")
                    except Exception as e:
                        st.warning(f"Could not parse/save {pdf.name}: {e}")
            
            # 2. Agent Execution
            agent = ScreeningAgent(
                run_id=st.session_state.run_id, 
                criteria=criteria, 
                scenario_type=scenario_type,
                simulate_failure=simulate_failure
            )
            results = agent.run(resumes)
            
            st.session_state.results = results
            st.session_state.metrics = agent.metrics
            
        st.success(f"Run {st.session_state.run_id} completed!")

    if st.session_state.results:
        st.subheader("📥 Export Reports")
        
        # Prepare CSV Data
        csv_data = []
        for r in st.session_state.results:
            row = {
                "Rank": r.rank,
                "Name": r.candidate_name,
                "Total Score": r.total_score,
                "Recommendation": r.recruiter_summary
            }
            for crit, score in r.scores.items():
                row[crit] = score
            csv_data.append(row)
        df_export = pd.DataFrame(csv_data)
        
        # CSV Download Button
        st.download_button(
            label="Download CSV Report",
            data=df_export.to_csv(index=False).encode('utf-8'),
            file_name=f"screening_report_{st.session_state.run_id}.csv",
            mime='text/csv',
            use_container_width=True
        )
        
        # JSON Download Button
        json_str = json.dumps([r.dict() for r in st.session_state.results], indent=4)
        st.download_button(
            label="Download JSON Report",
            data=json_str,
            file_name=f"screening_report_{st.session_state.run_id}.json",
            mime='application/json',
            use_container_width=True
        )

        # PDF Download Button
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Header
            pdf.set_fill_color(41, 128, 185) # Blue header
            pdf.rect(0, 0, 210, 40, 'F')
            pdf.set_text_color(255, 255, 255)
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 20, "RESUME SCREENING ASSISTANT", 0, 1, 'C')
            pdf.set_font("Arial", "", 10)
            pdf.cell(0, 10, f"Run ID: {st.session_state.run_id} | Date: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}", 0, 1, 'C')
            
            pdf.set_text_color(0, 0, 0)
            pdf.ln(10)

            # 1. Dashboard Overview (Visualization Section)
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "1. Dashboard Visual Analytics", 0, 1)
            pdf.ln(5)

            # --- Embedding Radar Chart (Using Matplotlib for robustness) ---
            try:
                top_3 = st.session_state.results[:3]
                categories = list(top_3[0].scores.keys())
                N = len(categories)
                angles = [n / float(N) * 2 * np.pi for n in range(N)]
                angles += angles[:1]

                fig, ax = plt.subplots(figsize=(8, 6), subplot_kw=dict(polar=True))
                plt.xticks(angles[:-1], categories, color='grey', size=8)
                
                for r in top_3:
                    values = list(r.scores.values())
                    values += values[:1]
                    ax.plot(angles, values, linewidth=1, linestyle='solid', label=r.candidate_name)
                    ax.fill(angles, values, alpha=0.1)
                
                plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1))
                plt.title("Top 3 Comparison", size=15, color='blue', y=1.1)
                
                radar_path = "temp_radar.png"
                plt.savefig(radar_path, bbox_inches='tight')
                plt.close()
                
                pdf.image(radar_path, x=10, w=190)
                os.remove(radar_path)
            except Exception as e:
                pdf.cell(0, 10, f"(Note: Radar skipped: {e})", 0, 1)

            # --- Embedding Bubble Chart (Using Matplotlib) ---
            try:
                df_b = pd.DataFrame([{"Rank": r.rank, "Score": r.total_score} for r in st.session_state.results])
                plt.figure(figsize=(10, 4))
                plt.scatter(df_b['Rank'], df_b['Score'], s=df_b['Score']*5, c=df_b['Score'], cmap='viridis', alpha=0.6)
                plt.title("Batch Match Strength Overview")
                plt.xlabel("Rank")
                plt.ylabel("Score")
                plt.grid(True, linestyle='--', alpha=0.7)
                
                bubble_path = "temp_bubble.png"
                plt.savefig(bubble_path, bbox_inches='tight')
                plt.close()
                
                pdf.image(bubble_path, x=10, w=190)
                os.remove(bubble_path)
            except Exception as e:
                pdf.cell(0, 10, f"(Note: Bubble skipped: {e})", 0, 1)

            pdf.ln(10)

            # 2. Executive Summary
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "2. Executive Summary", 0, 1)
            pdf.set_font("Arial", "", 11)
            pdf.cell(0, 7, f"Total Resumes Processed: {st.session_state.metrics.get('total_processed')}", 0, 1)
            pdf.cell(0, 7, f"Average Agent Score: {st.session_state.metrics.get('avg_score')}%", 0, 1)
            pdf.cell(0, 7, f"Baseline Score: {st.session_state.metrics.get('baseline_score')}%", 0, 1)
            pdf.ln(5)

            # 3. Candidate Deep-Dive
            pdf.add_page()
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, "3. Candidate Deep-Dive (Top 5)", 0, 1)
            pdf.ln(5)

            for r in st.session_state.results[:5]:
                # Sanitize text for PDF (Remove non-latin-1 characters like en-dash, em-dash, etc.)
                safe_name = r.candidate_name.upper().encode('latin-1', 'replace').decode('latin-1')
                safe_summary = r.recruiter_summary.encode('latin-1', 'replace').decode('latin-1')
                
                pdf.set_fill_color(240, 240, 240)
                pdf.set_font("Arial", "B", 12)
                pdf.cell(0, 10, f" RANK {r.rank}: {safe_name} ", 1, 1, 'L', True)
                
                pdf.set_font("Arial", "B", 10)
                pdf.cell(100, 8, "Criterion", 1)
                pdf.cell(40, 8, "Score", 1, 1, 'C')
                pdf.set_font("Arial", "", 10)
                for crit, score in r.scores.items():
                    pdf.cell(100, 8, crit, 1)
                    pdf.cell(40, 8, f"{score}%", 1, 1, 'C')
                
                pdf.set_font("Arial", "B", 10)
                pdf.cell(100, 8, "OVERALL MATCH SCORE", 1)
                pdf.cell(40, 8, f"{r.total_score}%", 1, 1, 'C')
                
                # --- Embedding Individual Chart (Using Matplotlib) ---
                try:
                    labels = list(r.scores.keys())
                    values = list(r.scores.values())
                    fig, ax = plt.subplots(figsize=(3, 3))
                    ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, wedgeprops=dict(width=0.4))
                    
                    donut_path = f"temp_donut_{r.resume_id}.png"
                    plt.savefig(donut_path, bbox_inches='tight')
                    plt.close()
                    
                    pdf.image(donut_path, x=145, y=pdf.get_y()-45, w=50) # Position next to scores
                    os.remove(donut_path)
                    
                    # Mention skills in PDF
                    pdf.set_font("Arial", "I", 8)
                    pdf.set_xy(145, pdf.get_y() + 5)
                    pdf.cell(50, 5, f"Metrics: {', '.join(labels)}", 0, 1, 'C')
                    
                    if r.matched_skills:
                        pdf.set_xy(145, pdf.get_y())
                        pdf.set_font("Arial", "B", 7)
                        pdf.multi_cell(50, 4, f"Matched: {', '.join(r.matched_skills)}", 0, 'C')
                    
                    pdf.set_xy(10, pdf.get_y() + 5)
                except:
                    pass

                pdf.ln(2)
                pdf.set_font("Arial", "I", 10)
                pdf.multi_cell(0, 5, f"Agent Advice: {safe_summary}")
                
                pdf.set_font("Arial", "B", 10)
                pdf.cell(0, 8, "Tailored Interview Questions:", 0, 1)
                pdf.set_font("Arial", "", 10)
                for q in (r.interview_questions or []):
                    safe_q = q.encode('latin-1', 'replace').decode('latin-1')
                    pdf.cell(0, 6, f"- {safe_q}", 0, 1)
                pdf.ln(10)

            pdf_output = pdf.output()
            # Ensure output is bytes
            if isinstance(pdf_output, bytearray):
                pdf_data = bytes(pdf_output)
            elif isinstance(pdf_output, str):
                pdf_data = pdf_output.encode('latin1')
            else:
                pdf_data = pdf_output

            st.download_button(
                label="Download Visual PDF Report",
                data=pdf_data,
                file_name=f"visual_report_{st.session_state.run_id}.pdf",
                mime='application/pdf',
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Error generating PDF: {e}")

        # Word (DOCX) Download Button
        try:
            doc = Document()
            doc.add_heading('RESUME SCREENING ASSISTANT REPORT', 0)
            doc.add_paragraph(f"Run ID: {st.session_state.run_id}")
            doc.add_paragraph(f"Date: {pd.Timestamp.now().strftime('%Y-%m-%d')}")

            doc.add_heading('1. Executive Summary', level=1)
            doc.add_paragraph(f"Total Processed: {st.session_state.metrics.get('total_processed')}")
            doc.add_paragraph(f"Batch Avg Score: {st.session_state.metrics.get('avg_score')}%")

            # --- Embedding Visualizations in Word (Using Matplotlib for robustness) ---
            try:
                doc.add_heading('2. Batch Visual Analytics', level=1)
                
                # Radar Chart
                top_3 = st.session_state.results[:3]
                categories = list(top_3[0].scores.keys())
                N = len(categories)
                angles = [n / float(N) * 2 * np.pi for n in range(N)]
                angles += angles[:1]

                fig, ax = plt.subplots(figsize=(6, 5), subplot_kw=dict(polar=True))
                plt.xticks(angles[:-1], categories, size=8)
                for r in top_3:
                    values = list(r.scores.values())
                    values += values[:1]
                    ax.plot(angles, values, linewidth=1, label=r.candidate_name)
                    ax.fill(angles, values, alpha=0.1)
                plt.legend(loc='upper right', bbox_to_anchor=(0.1, 0.1), fontsize='small')
                
                radar_io = BytesIO()
                plt.savefig(radar_io, format='png', bbox_inches='tight')
                plt.close()
                radar_io.seek(0)
                doc.add_picture(radar_io, width=docx.shared.Inches(5))
                
                # Mention skills in Word
                skills_text = " | ".join(categories)
                p = doc.add_paragraph()
                p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"Analyzed Criteria: {skills_text}")
                run.bold = True
                run.font.size = docx.shared.Pt(10)
                
                # Bubble Chart
                plt.figure(figsize=(8, 3))
                df_b = pd.DataFrame([{"Rank": r.rank, "Score": r.total_score} for r in st.session_state.results])
                plt.scatter(df_b['Rank'], df_b['Score'], s=df_b['Score']*4, c=df_b['Score'], cmap='viridis', alpha=0.5)
                plt.title("Batch Performance Overview")
                
                bubble_io = BytesIO()
                plt.savefig(bubble_io, format='png', bbox_inches='tight')
                plt.close()
                bubble_io.seek(0)
                doc.add_picture(bubble_io, width=docx.shared.Inches(5))
                
            except Exception as e:
                doc.add_paragraph(f"(Visualizations skipped: {str(e)})")

            doc.add_heading('3. Candidate Breakdown (Top 10)', level=1)
            for r in st.session_state.results[:10]:
                doc.add_heading(f"Rank {r.rank}: {r.candidate_name}", level=2)
                doc.add_paragraph(f"Overall Score: {r.total_score}%")
                doc.add_paragraph(f"Recommendation: {r.recruiter_summary}")
                
                # Mini Chart for Candidate (Using Matplotlib Pie)
                try:
                    labels = list(r.scores.keys())
                    values = list(r.scores.values())
                    fig, ax = plt.subplots(figsize=(3, 3))
                    ax.pie(values, labels=labels, autopct='%1.1f%%', startangle=90, wedgeprops=dict(width=0.4))
                    ax.set_title("Score Breakdown", size=10)
                    
                    donut_io = BytesIO()
                    plt.savefig(donut_io, format='png', bbox_inches='tight')
                    plt.close()
                    donut_io.seek(0)
                    doc.add_picture(donut_io, width=docx.shared.Inches(2.5))
                    
                    # Skills for this candidate
                    p_skills = doc.add_paragraph()
                    p_skills.add_run(f"Analyzed Categories: {', '.join(labels)}").italic = True
                    
                    if r.matched_skills:
                        p_m = doc.add_paragraph()
                        p_m.add_run(f"Matched Skills: {', '.join(r.matched_skills)}").bold = True
                except:
                    pass

                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Criterion'
                hdr_cells[1].text = 'Score'
                for crit, score in r.scores.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = crit
                    row_cells[1].text = f"{score}%"
                
                doc.add_paragraph("Interview Questions:")
                for q in (r.interview_questions or []):
                    doc.add_paragraph(f"{q}", style='List Bullet')

            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            st.download_button(
                label="Download Visual Word Report",
                data=doc_io,
                file_name=f"visual_report_{st.session_state.run_id}.docx",
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Error generating Word report: {e}")

    if st.button("🔄 Reset App", use_container_width=True):
        st.session_state.run_id = None
        st.session_state.results = []
        st.session_state.metrics = {}
        if os.path.exists("runs.jsonl"):
            os.remove("runs.jsonl")
        st.rerun()

    if st.session_state.metrics:
        st.subheader("📊 Screening Analytics")
        m = st.session_state.metrics
        st.info(f"🏗️ **Infrastructure:** {m.get('infrastructure', 'N/A')}")
        
        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.metric("Processed", m.get("total_processed", 0))
        
        avg_score = m.get("avg_score", 0)
        baseline = m.get("baseline_score", 0)
        delta = round(avg_score - baseline, 2) if baseline > 0 else 0
        
        col_m2.metric("Agent Score", f"{avg_score}%", delta=f"{delta}% vs baseline")
        col_m3.metric("Execution Time", f"{m.get('total_time', 0)}s")
        
        if "error" in st.session_state.metrics:
            st.error(f"Run failed: {st.session_state.metrics['error']}")

with col2:
    st.subheader("🕵️ Agent Observability")
    
    if st.session_state.run_id:
        st.info(f"**Active Agents:** `TechnicalAgent`, `SoftSkillsAgent`, `AuditorAgent`, `RecruiterAgent` | **Run ID:** `{st.session_state.run_id}`")
        tabs = st.tabs(["State Transitions", "Tool Logs", "Transcript", "Analytics", "Final Reports", "Batch Insights", "Run History", "📖 Agent Guide"])
        
        # Load logs for this session
        logs = []
        if os.path.exists("runs.jsonl"):
            with open("runs.jsonl", "r") as f:
                for line in f:
                    entry = json.loads(line)
                    if entry.get("run_id") == st.session_state.run_id:
                        logs.append(entry)
        
        with tabs[0]:
            st.subheader("🔄 State Transitions")
            transitions = [l for l in logs if l.get("next_state") is not None]
            if transitions:
                df_trans = pd.DataFrame(transitions)
                st.dataframe(df_trans[["timestamp", "previous_state", "event", "next_state"]], use_container_width=True)
            else:
                st.info("No transitions logged yet.")

        with tabs[1]:
            st.subheader("🛠️ Tool Execution Logs")
            tool_calls = [l for l in logs if l.get("tool_name") is not None]
            if tool_calls:
                df_tools = pd.DataFrame(tool_calls)
                st.dataframe(df_tools[["timestamp", "tool_name", "tool_input", "tool_output"]], use_container_width=True)
            else:
                st.info("No tool calls logged yet.")
                
        with tabs[2]:
            st.subheader("🤖 Agent Reasoning & Transcript")
            if logs:
                for entry in logs:
                    if entry.get("type") == "STATE_TRANSITION":
                        st.markdown(f"**State Changed:** `{entry['previous_state']} ➔ {entry['next_state']}`")
                    elif entry.get("type") == "TOOL_CALL":
                        with st.expander(f"🛠️ Tool: {entry['tool']}"):
                            st.write(f"**Thought:** *{entry.get('thought', 'N/A')}*" 
                                     if entry.get("thought") else "**Thought:** N/A")
                            st.json({"Input": entry["tool_input"], "Output": entry["tool_output"]})
                    elif entry.get("event"):
                        st.write(f"**[{entry['timestamp']}]** {entry['event']} ➔ `{entry['next_state']}`")
                    elif entry.get("tool_name"):
                        st.write(f"**[{entry['timestamp']}]** Called Agent: `{entry['tool_name']}`")
            else:
                st.info("No logs found for this run.")
        
        with tabs[3]:
            if st.session_state.results:
                col_a1, col_i2 = st.columns([1, 1])
                
                with col_a1:
                    st.subheader("Top Candidate Match")
                    top_cand = st.session_state.results[0]
                    fig_gauge = go.Figure(go.Indicator(
                        mode = "gauge+number",
                        value = top_cand.total_score,
                        domain = {'x': [0, 1], 'y': [0, 1]},
                        title = {'text': f"{top_cand.candidate_name}"},
                        gauge = {
                            'axis': {'range': [None, 100]},
                            'bar': {'color': "#2980b9"},
                            'steps' : [
                                {'range': [0, 50], 'color': "#ebedef"},
                                {'range': [50, 80], 'color': "#aed6f1"},
                                {'range': [80, 100], 'color': "#5dade2"}],
                            'threshold': {
                                'line': {'color': "red", 'width': 4},
                                'thickness': 0.75,
                                'value': 90}
                        }
                    ))
                    fig_gauge.update_layout(height=300, margin=dict(l=20, r=20, t=50, b=20), template="plotly_dark")
                    st.plotly_chart(fig_gauge, use_container_width=True)

                with col_i2:
                    st.subheader("Top 3 Skills Radar")
                    top_3 = st.session_state.results[:3]
                    fig_radar = go.Figure()
                    for r in top_3:
                        categories = list(r.scores.keys())
                        values = list(r.scores.values())
                        fig_radar.add_trace(go.Scatterpolar(r=values, theta=categories, fill='toself', name=r.candidate_name))
                    fig_radar.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 100])), height=300, showlegend=True, template="plotly_dark", margin=dict(l=20, r=20, t=50, b=20))
                    st.plotly_chart(fig_radar, use_container_width=True)
                    
                    # Mention analyzed criteria names below the radar chart
                    if top_3:
                        skills_list = " | ".join([f"<b>{c}</b>" for c in list(top_3[0].scores.keys())])
                        matched_list = " | ".join([f"<span style='color:#5dade2;'>{s}</span>" for s in (top_3[0].matched_skills or [])])
                        
                        st.markdown(f"""
                        <div style='text-align: center; background-color: #1e1e1e; padding: 15px; border-radius: 10px; border: 1px solid #333;'>
                            <div style='margin-bottom: 8px;'>🎯 <b>Top Analyzed Categories:</b> {skills_list}</div>
                            <div style='font-size: 0.9em; border-top: 1px solid #444; padding-top: 8px;'>✅ <b>Specific Skills Matched:</b> {matched_list if matched_list else 'N/A'}</div>
                        </div>
                        """, unsafe_allow_html=True)

                st.subheader("Batch Distribution (Bubble View)")
                # Bubble chart - circular visualization of all candidates
                df_bubble = pd.DataFrame([
                    {"Name": r.candidate_name, "Score": r.total_score, "Rank": r.rank} 
                    for r in st.session_state.results
                ])
                fig_bubble = px.scatter(df_bubble, x="Rank", y="Score", size="Score", color="Score",
                                        hover_name="Name", title="Batch Overview (Size = Match Strength)",
                                        color_continuous_scale=px.colors.sequential.Viridis)
                fig_bubble.update_layout(template="plotly_dark", height=400)
                st.plotly_chart(fig_bubble, use_container_width=True)
            else:
                st.info("Run screening to see visual analytics.")
        
        with tabs[4]:
            if st.session_state.results:
                st.subheader("Top 5 Candidate Deep-Dive")
                for r in st.session_state.results[:5]:
                    with st.expander(f"Rank {r.rank}: {r.candidate_name} (Score: {r.total_score})"):
                        col_d1, col_d2 = st.columns([2, 1])
                        with col_d1:
                            st.write(f"**Agent Recommendation:** {r.recruiter_summary}")
                            st.write("**Tailored Interview Questions:**")
                            for q in (r.interview_questions or []):
                                st.write(f"- {q}")
                        
                        with col_d2:
                            # Mini Donut Chart for individual candidate breakdown
                            fig_donut = go.Figure(data=[go.Pie(labels=list(r.scores.keys()), 
                                                            values=list(r.scores.values()), 
                                                            hole=.6)])
                            fig_donut.update_layout(showlegend=False, height=200, margin=dict(l=0, r=0, t=0, b=0), template="plotly_dark")
                            st.plotly_chart(fig_donut, use_container_width=True, key=f"donut_{r.resume_id}_{st.session_state.run_id}")
                            
                            # Explicitly mention top skills for this candidate
                            skills_summary = ", ".join(list(r.scores.keys()))
                            st.caption(f"📊 **Metrics Analyzed:** {skills_summary}")
                            
                            # Show matched specific skills for this candidate
                            if r.matched_skills:
                                st.markdown(f"✅ **Matched Skills:** {', '.join(r.matched_skills)}")
                            else:
                                st.caption("No specific skills extracted yet.")
            else:
                st.info("Detailed reports will be generated after a run.")

        with tabs[5]:
            st.subheader("📈 Batch Insights & Skills Gap")
            if st.session_state.results:
                col_i1, col_i2 = st.columns(2)
                with col_i1:
                    st.success("**Strengths Found in Batch**")
                    st.write("- High technical alignment in core stack")
                    st.write("- Strong educational background across top 10")
                with col_i2:
                    st.warning("**Potential Gaps**")
                    st.write("- Lack of leadership experience in junior roles")
                    st.write("- Niche skills (Cloud/DevOps) are rare")
                
                # Dynamic Correlation
                df_corr = pd.DataFrame([{"Exp": r.scores.get("Experience Match", 0), "Total": r.total_score} for r in st.session_state.results])
                fig_corr = px.scatter(df_corr, x="Exp", y="Total", title="Experience vs Total Score Correlation")
                st.plotly_chart(fig_corr, use_container_width=True)
            else:
                st.info("Run screening for batch insights.")

        with tabs[6]:
            st.subheader("📜 Run History & Audit Trail")
            if os.path.exists("runs.jsonl"):
                history = {}
                with open("runs.jsonl", "r") as f:
                    for line in f:
                        e = json.loads(line)
                        rid = e.get("run_id")
                        if rid not in history: history[rid] = {"start": e["timestamp"], "status": "In Progress"}
                        if e.get("next_state") == "COMPLETED": history[rid]["status"] = "✅ Success"
                
                df_hist = pd.DataFrame([{"Run ID": k, "Time": v["start"], "Status": v["status"]} for k, v in history.items()]).sort_values("Time", ascending=False)
                st.dataframe(df_hist, use_container_width=True, hide_index=True)
            else:
                st.info("No audit logs found.")

        with tabs[7]:
            st.subheader("📖 Understanding the Multi-Agent Brain")
            st.markdown("""
            This system uses a **Decentralized Multi-Agent Architecture** where each agent has a specific responsibility:
            
            1. **💻 Technical Agent**: 
               - Uses **LLM Deep Reasoning** to analyze resume content.
               - Performs **Semantic Matching** (not just keywords) against job requirements.
               - Generates contextual reasoning for its technical score.
            
            2. **🤝 Soft Skills Agent**:
               - **Role-Aware Analysis**: Adjusts its criteria based on the job scenario (e.g., DevOps vs Frontend).
               - Looks for leadership, communication, and adaptability patterns in the narrative.
            
            3. **⚖️ Auditor Agent**:
               - Acts as the **Final Judge**.
               - Weighs Technical Match (85%) against Soft Skills (15%).
               - Calculates the **Final Consolidated Score**.
            
            4. **📝 Recruiter Agent**:
               - The **Reporting Specialist**.
               - Generates human-readable summaries and **tailored interview questions** based on candidate gaps.
            
            ---
            **Tech Stack:** Streamlit (UI), Python (Agents), RAG Simulation (Vector Store), JSONL (Observability Logs).
            """)
    else:
        st.info("Start a run to see agent activity.")

st.markdown("--- ")
col_res, col_feed = st.columns([2, 1])

with col_res:
    st.subheader("🏆 Ranked Candidates")
    if st.session_state.results:
        data = []
        for r in st.session_state.results:
            row = {
                "Rank": r.rank,
                "Name": r.candidate_name,
                "Total Score": r.total_score
            }
            for crit_name, score in r.scores.items():
                row[crit_name] = score
            data.append(row)
        
        df_results = pd.DataFrame(data)
        st.dataframe(df_results, use_container_width=True, hide_index=True)
    else:
        st.write("No results yet.")

with col_feed:
    st.subheader("✍️ Human-in-the-loop")
    if st.session_state.results:
        selected_candidate = st.selectbox("Adjust Score for:", [r.candidate_name for r in st.session_state.results])
        # 0 to 100 points adjustment (Override or Offset)
        adjustment = st.slider("Manual Score Adjustment (0-100)", 0, 100, 50, help="Override or add extra weight to this candidate's score.")
        
        if st.button("Apply Feedback"):
            # Logic: Add a percentage of the manual adjustment to the total score
            for r in st.session_state.results:
                if r.candidate_name == selected_candidate:
                    # We store the base score if not already stored to allow multiple adjustments
                    base_key = f"base_score_{r.resume_id}"
                    if base_key not in st.session_state:
                        st.session_state[base_key] = r.total_score
                    
                    # Apply adjustment: Base Score (70%) + Human Feedback (30%)
                    r.total_score = round((st.session_state[base_key] * 0.7) + (adjustment * 0.3), 2)
                    st.session_state.results = sorted(st.session_state.results, key=lambda x: x.total_score, reverse=True)
                    for i, res in enumerate(st.session_state.results): res.rank = i + 1
                    st.success(f"Adjusted {selected_candidate}'s score based on human intuition!")
                    st.rerun()
    else:
        st.info("Run screening to enable feedback.")

# Footer
st.markdown("---")
st.markdown(
        """
        <div style='text-align: center; color: grey;'>
            <p>Resume Screening Assistant | Built with Multi-Agent RAG Architecture</p>
            <p><i>Standard Compliance: JSONL Logging | Seed-based Reproducibility | State Machine Observability</i></p>
        </div>
        """,
        unsafe_allow_html=True
    )
